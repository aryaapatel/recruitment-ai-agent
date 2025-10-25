# # # app/parse_utils.py
# # import io
# # from typing import Union
# # from PyPDF2 import PdfReader
# # import docx

# # def extract_text_from_file(file: Union[io.BytesIO, str]):
# #     filename = getattr(file, "filename", str(file))
# #     if filename.endswith(".pdf"):
# #         reader = PdfReader(file)
# #         text = "\n".join(page.extract_text() or "" for page in reader.pages)
# #         return text
# #     elif filename.endswith(".docx"):
# #         doc = docx.Document(file)
# #         text = "\n".join([p.text for p in doc.paragraphs])
# #         return text
# #     elif filename.endswith(".txt"):
# #         if hasattr(file, "read"):
# #             return file.read().decode("utf-8")
# #         with open(file, "r") as f:
# #             return f.read()
# #     else:
# #         return ""
# # app/parse_utils.py
# # app/parse_utils.py
# import io
# from typing import Union
# from PyPDF2 import PdfReader
# import docx
# from fastapi import UploadFile

# # app/parse_utils.py
# import io
# from typing import Union
# from PyPDF2 import PdfReader
# import docx
# from fastapi import UploadFile

# async def extract_text_from_file(file: Union[UploadFile, str]):
#     filename = getattr(file, "filename", str(file))

#     # Step 1: Read content into bytes
#     if isinstance(file, UploadFile):
#         content = await file.read()  # await the coroutine
#         file_bytes = io.BytesIO(content)  # wrap in BytesIO
#     else:
#         # Already a path or file-like object
#         if isinstance(file, str):
#             with open(file, "rb") as f:
#                 file_bytes = io.BytesIO(f.read())
#         else:
#             file_bytes = file

#     # Step 2: Extract text based on extension
#     if filename.lower().endswith(".pdf"):
#         reader = PdfReader(file_bytes)  # now safe
#         text = "\n".join(page.extract_text() or "" for page in reader.pages)
#         return text

#     elif filename.lower().endswith(".docx"):
#         doc = docx.Document(file_bytes)
#         text = "\n".join([p.text for p in doc.paragraphs])
#         return text

#     elif filename.lower().endswith(".txt"):
#         return file_bytes.getvalue().decode("utf-8")

#     else:
#         return ""

# app/parse_utils.py
import io
from PyPDF2 import PdfReader
import docx

def extract_text_from_file(file_bytes: io.BytesIO, filename: str) -> str:
    """
    Extracts text from a file given as a byte stream and its filename.
    This function is synchronous and framework-agnostic.
    """
    try:
        if filename.lower().endswith(".pdf"):
            reader = PdfReader(file_bytes)
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        elif filename.lower().endswith(".docx"):
            doc = docx.Document(file_bytes)
            return "\n".join([p.text for p in doc.paragraphs])

        elif filename.lower().endswith(".txt"):
            return file_bytes.getvalue().decode("utf-8")

        else:
            return f"Unsupported file type: {filename}"
            
    except Exception as e:
        print(f"Error processing file {filename}: {e}")
        return f"Error extracting text from {filename}."